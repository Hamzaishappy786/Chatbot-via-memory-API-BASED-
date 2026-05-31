import base64
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from app.ingestion.parsers.base import BaseParser, ParsedContent


class DOCXParser(BaseParser):
    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]

    def parse(self, file_path: str) -> ParsedContent:
        doc = Document(file_path)
        content = ParsedContent()

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                header = rows[0]
                separator = " | ".join(["---"] * len(table.columns))
                table_md = "\n".join([header, separator] + rows[1:])
                paragraphs.append(f"\n{table_md}\n")

        if paragraphs:
            content.text_blocks.append({
                "text": "\n\n".join(paragraphs),
                "page": 1,
            })

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    img_b64 = base64.b64encode(img_bytes).decode()
                    content.images.append({
                        "image_base64": img_b64,
                        "page": 1,
                        "source": "embedded",
                    })
                except Exception:
                    continue

        return content
